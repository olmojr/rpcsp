import streamlit as st
import fitz  # PyMuPDF
import os
import unicodedata
import re
import time
import io
from reportlab.lib.styles import ParagraphStyle
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.pagesizes import A4
from reportlab.lib.units import inch
from reportlab.lib.enums import TA_JUSTIFY, TA_LEFT, TA_CENTER, TA_RIGHT
from reportlab.lib.colors import black
from reportlab.lib.fonts import addMapping
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from docx import Document

# Título de la aplicación
st.title("Resaltador de Palabras Clave")

# Descripción
st.write("""
Sube un archivo DOCX o pega el texto de cualquier documento y resalta las palabras clave en el PDF generado.
Puedes especificar si quieres coincidencia exacta o parcial para cada palabra.
""")

def get_formatted_text(paragraph):
    """
    Convierte un párrafo de Word a texto formateado con HTML para ReportLab
    """
    text_parts = []
    
    for run in paragraph.runs:
        text = run.text
        if text.strip():  # Solo procesar si hay texto
            # Inicio del formato
            format_start = ''
            format_end = ''
            
            # Aplicar negrita
            if run.bold:
                format_start += '<b>'
                format_end = '</b>' + format_end
            
            # Aplicar cursiva
            if run.italic:
                format_start += '<i>'
                format_end = '</i>' + format_end
            
            # Aplicar subrayado
            if run.underline:
                format_start += '<u>'
                format_end = '</u>' + format_end
                
            # Combinar el texto con sus formatos
            formatted_text = format_start + text + format_end
            text_parts.append(formatted_text)
    
    return ' '.join(text_parts)

def get_paragraph_alignment(paragraph):
    """
    Determina la alineación del párrafo
    """
    if paragraph.alignment == 1:
        return TA_CENTER
    elif paragraph.alignment == 2:
        return TA_RIGHT
    elif paragraph.alignment == 3:
        return TA_JUSTIFY
    else:
        return TA_LEFT

# Función para eliminar acentos
def remove_accents(input_str):
    nfkd_form = unicodedata.normalize('NFKD', input_str)
    return "".join([c for c in nfkd_form if not unicodedata.combining(c)])

# Función para verificar si un archivo existe y esperar si es necesario
def wait_for_file(file_path, timeout=30):
    start_time = time.time()
    while not os.path.exists(file_path):
        if time.time() - start_time > timeout:
            raise TimeoutError(f"Timeout esperando por el archivo: {file_path}")
        time.sleep(1)
    return True

def docx_to_pdf(docx_path, pdf_path):
    # Leer el documento DOCX
    doc = Document(docx_path)
    
    # Configurar el documento PDF
    pdf = SimpleDocTemplate(
        pdf_path,
        pagesize=A4,
        rightMargin=72,
        leftMargin=72,
        topMargin=72,
        bottomMargin=72
    )
    
    # Crear estilo base para el texto
    base_style = ParagraphStyle(
        'BaseStyle',
        fontSize=11,
        leading=14,
        spaceBefore=0,
        spaceAfter=12,
        allowWidows=0,
        allowOrphans=0
    )
    
    # Lista para almacenar los elementos del documento
    story = []
    
    # Procesar cada párrafo
    for para in doc.paragraphs:
        if para.text.strip():  # Solo procesar párrafos con texto
            # Crear un estilo específico para este párrafo
            para_style = ParagraphStyle(
                'ParaStyle',
                parent=base_style,
                alignment=get_paragraph_alignment(para)
            )
            
            # Obtener el texto formateado con HTML
            formatted_text = get_formatted_text(para)
            
            # Crear el párrafo con el estilo y el texto formateado
            p = Paragraph(formatted_text, para_style)
            story.append(p)
            story.append(Spacer(1, 6))
    
    # Construir el PDF
    try:
        pdf.build(story)
    except Exception as e:
        st.error(f"Error al generar el PDF: {str(e)}")
        raise

# Función para resaltar palabras exactas en PDF
def highlight_words_in_pdf(input_pdf, output_pdf, words_to_highlight, exact_match_words):
    if not os.path.exists(input_pdf):
        raise FileNotFoundError(f"No se encuentra el archivo PDF de entrada: {input_pdf}")
        
    doc = fitz.open(input_pdf)
    palabras_no_encontradas = set(words_to_highlight)

    # Definir los caracteres de puntuación que queremos considerar
    punctuation = r'[.,;:!¡?¿"\')\]}]'
    
    for page_num in range(len(doc)):
        page = doc.load_page(page_num)
        text = page.get_text("text")
        text_normalized = remove_accents(text).lower()

        for word_to_highlight in words_to_highlight:
            word_normalized = remove_accents(word_to_highlight).lower()
            
            if word_to_highlight in exact_match_words:
                # Búsqueda exacta modificada para incluir puntuación
                regex_pattern = r"\b" + re.escape(word_normalized) + r"(?:\s|" + punctuation + r"|$)"
            else:
                # Búsqueda parcial (sin cambios)
                regex_pattern = re.escape(word_normalized)
                
            regex = re.compile(regex_pattern, re.IGNORECASE)
            
            # Buscar coincidencias
            if regex.search(text_normalized):
                palabras_no_encontradas.discard(word_to_highlight)
                
            # Resaltar coincidencias
            areas = page.search_for(word_to_highlight)
            
            if word_to_highlight in exact_match_words:
                # Para palabras exactas, verificar que cada coincidencia sea una palabra completa
                for area in areas:
                    rect = area
                    expanded_rect = fitz.Rect(
                        rect.x0 - 2,
                        rect.y0 - 2,
                        rect.x1 + 5,
                        rect.y1 + 2
                    )
                    
                    words_around = page.get_text("words", clip=expanded_rect)
                    
                    # Verificar si es una palabra completa (incluyendo puntuación)
                    for word_info in words_around:
                        word_text = word_info[4]
                        # Eliminar la puntuación al final de la palabra para la comparación
                        word_text_clean = re.sub(punctuation + r'$', '', word_text)
                        if remove_accents(word_text_clean).lower() == word_normalized:
                            page.add_highlight_annot(area)
                            break
            else:
                # Para coincidencias parciales, resaltar todas las ocurrencias
                for area in areas:
                    page.add_highlight_annot(area)

    doc.save(output_pdf)
    doc.close()
    return list(palabras_no_encontradas)

# Función para limpiar archivos temporales
def cleanup_temp_files(files):
    for file in files:
        try:
            if os.path.exists(file):
                os.remove(file)
        except Exception as e:
            st.warning(f"No se pudo eliminar el archivo temporal {file}: {str(e)}")

# Selector de método de entrada
input_method = st.radio(
    "Selecciona el método de entrada:",
    ["Subir archivo DOCX", "Pegar texto"],
    help="Elige si quieres subir un archivo DOCX o pegar directamente el texto"
)

# Variable para almacenar el contenido a procesar
content_to_process = None

if input_method == "Subir archivo DOCX":
    # Subir archivo DOCX
    uploaded_file = st.file_uploader("Sube un archivo DOCX", type=["docx"])
    if uploaded_file is not None:
        content_to_process = "file"
else:
    # Área de texto para pegar contenido
    pasted_text = st.text_area(
        "Pega aquí el contenido del documento:",
        height=300,
        help="Pega el texto de documentos DOC, ODT u otros formatos"
    )
    
    # Botón para procesar el texto pegado
    if st.button("Procesar texto", 
                 help="Haz clic para procesar el texto pegado",
                 disabled=not pasted_text.strip()):
        if pasted_text.strip():
            content_to_process = "text"
        else:
            st.warning("Por favor, pega algún texto antes de procesar.")

# Campo para ingresar palabras clave con selector de coincidencia exacta
st.write("Ingresa las palabras a resaltar y selecciona el tipo de coincidencia para cada una:")

# Crear dos columnas
col1, col2 = st.columns([3, 1])

# Palabras por defecto
default_words = "Metodologías activas, ODS, Situación de aprendizaje, XXI, Competencias clave, Objetivos de etapa, Atención a la diversidad, Diferencias individuales, DUA, Competencias específicas, Criterios de evaluación, Bloque de contenidos, Reto, Sesiones, Producto final, Coordinación docente, Centro, Familia, Competencia digital"
default_exact_matches = ["DUA", "Reto"]

# En la primera columna, el input de palabras
with col1:
    words_input = st.text_input(
        "Palabras (separadas por comas)",
        default_words
    )

# En la segunda columna, selector múltiple para coincidencias exactas
with col2:
    words_list = [word.strip() for word in words_input.split(",") if word.strip()]
    exact_match_words = st.multiselect(
        "Coincidencia exacta",
        options=words_list,
        default=default_exact_matches,
        help="Selecciona las palabras que requieren coincidencia exacta"
    )

# Procesar el contenido si existe
if content_to_process:
    # Crear nombres de archivos temporales con rutas absolutas
    temp_dir = os.path.abspath(os.path.dirname(__file__))
    temp_docx = os.path.join(temp_dir, "temp_input.docx")
    temp_pdf = os.path.join(temp_dir, "temp_output.pdf")
    output_pdf = os.path.join(temp_dir, "documento_revisado.pdf")

    # Lista de archivos temporales para limpieza
    temp_files = [temp_docx, temp_pdf, output_pdf]

    try:
        # Limpiar archivos temporales anteriores
        cleanup_temp_files(temp_files)

        if content_to_process == "file":
            # Guardar el archivo subido
            with open(temp_docx, "wb") as f:
                f.write(uploaded_file.getbuffer())
        else:
            # Crear documento DOCX desde el texto pegado
            doc = Document()
            doc.add_paragraph(pasted_text)
            doc.save(temp_docx)

        # Convertir DOCX a PDF usando la nueva función
        with st.spinner("Convirtiendo a PDF..."):
            docx_to_pdf(temp_docx, temp_pdf)
            
            if not wait_for_file(temp_pdf):
                raise FileNotFoundError("No se pudo generar el archivo PDF")

        # Resaltar palabras en PDF
        with st.spinner("Resaltando palabras en PDF..."):
            palabras_no_encontradas = highlight_words_in_pdf(
                temp_pdf, 
                output_pdf, 
                words_list,
                exact_match_words
            )

        if not os.path.exists(output_pdf):
            raise FileNotFoundError("No se pudo generar el archivo PDF final")

        # Mostrar el PDF resultante
        st.success("¡Proceso completado!")
        st.write("Descarga tu archivo PDF con las palabras resaltadas:")

        with open(output_pdf, "rb") as f:
            st.download_button(
                label="Descargar PDF",
                data=f,
                file_name="documento_revisado.pdf",
                mime="application/pdf"
            )

        # Mostrar palabras no encontradas
        if palabras_no_encontradas:
            st.warning("⚠️ Las siguientes palabras no se encontraron en el documento:")
            for palabra in palabras_no_encontradas:
                st.write(f"- {palabra}")
        else:
            st.info("Todas las palabras se encontraron y resaltaron correctamente.")

    except Exception as e:
        st.error(f"Error en el proceso: {str(e)}")
        st.error("Detalles del error para depuración:")
        st.error(f"- Directorio actual: {os.getcwd()}")
        st.error(f"- Archivo DOCX existe: {os.path.exists(temp_docx)}")
        st.error(f"- Archivo PDF temporal existe: {os.path.exists(temp_pdf)}")
        
    finally:
        # Limpiar archivos temporales
        cleanup_temp_files(temp_files)
